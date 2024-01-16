import random

from aiogram import Dispatcher, types
from aiogram.dispatcher import FSMContext
from handlers.states import SG
from aiogram.utils import exceptions
from models.adapters.base_adapter import BaseAdapter
from models.adapters.admin_panel_adapter import AdminPanelAdapter
from models.adapters.user_adapter import UserAdapter
from models.adapters.form_adapter import FormAdapter
from pagination.MagicPaginator import ButtonPagination
from aiogram.types import InlineKeyboardButton, InlineKeyboardMarkup
from docx import Document
import pathlib

from apscheduler.schedulers.asyncio import AsyncIOScheduler


async def get_contract_btn(call: types.CallbackQuery, state: FSMContext):
    client_id = call.from_user.id
    admin_panel_adapter = AdminPanelAdapter()
    text, keyboard = await admin_panel_adapter.get_user_contract_if_ready(client_id)
    if text.split("#")[0] == "1":
        admin_panel_adapter = AdminPanelAdapter()
        client_fio = await admin_panel_adapter.get_client_fio_by_id(client_id)
        contract_path = str(pathlib.Path.cwd()) + f"/files/contracts/Договор | {client_fio}.docx"
        with open(contract_path, "rb") as file:
            await call.message.bot.send_document(chat_id=client_id, document=file)

    await call.message.edit_text(text=text.split("#")[1], reply_markup=keyboard)


async def choose_client_to_create_contract(call: types.CallbackQuery, state: FSMContext):
    await state.update_data(call=call)
    admin_panel_adapter = AdminPanelAdapter()
    keyboard_data = await admin_panel_adapter.get_contract_waiting_users_for_partner_as_buttons(call.from_user.id)
    paginator = ButtonPagination(button_data=keyboard_data, amount_elements=5)
    custom_keyboard = [InlineKeyboardButton(text='◀︎', callback_data='back_to_partner_menu_btn')]
    await state.update_data(paginator=paginator)
    try:
        keyboard = paginator.pagination(custom_keyboard=custom_keyboard)
        await call.message.edit_text(text="<b>Выберите клиента:</b>", reply_markup=keyboard)
        await SG.clients_waiting_contract_pagination.set()
    except KeyError:
        await call.message.edit_text(text="<b>Клиентов на этапе формирования договора пока нет!</b>\n\n"
                                          "<i>Обратите внимение, что для формирования договора, клиент должен"
                                          " быть закреплен за вами и у него должна быть заполнена анкета.</i>",
                                     reply_markup=InlineKeyboardMarkup().add(InlineKeyboardButton('◀️   Назад    ', callback_data='back_to_partner_menu_btn')))














async def enter_birth_date(call: types.CallbackQuery, state: FSMContext):
    client_id = call.data.split("#")[-1]
    await state.update_data(client_id=client_id)
    form_adapter = FormAdapter()
    user_form = await form_adapter.get_form(client_id)
    await state.update_data(fio=user_form.user_fio)
    await state.update_data(phone_number=user_form.user_number)
    await call.message.edit_text(text="<b>Введите дату рождения клиента:</b>")
    await SG.enter_passport_data.set()


async def enter_passport_data(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(birth_date=message.text)
    await call.message.edit_text(text="<b>Введите серию и номер паспорта:</b>")
    await SG.enter_date.set()


async def enter_date(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(passport_data=message.text)
    await call.message.edit_text(text="<b>Введите дату договора:</b>")
    await SG.enter_passport_issuer.set()


async def enter_passport_issuer(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(date=message.text)
    await call.message.edit_text(text="<b>Введите кем выдан паспорт:</b>")
    await SG.enter_passport_issue_date.set()



async def enter_passport_issue_date(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(issuer=message.text)
    await call.message.edit_text(text="<b>Введите дату выдачи паспорта:</b>")
    await SG.enter_passport_issue_code.set()



async def enter_passport_issue_code(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(pass_issue_date=message.text)
    await call.message.edit_text(text="<b>Введите код подразделения паспорта:</b>")
    await SG.enter_credit_sum.set()



async def enter_credit_sum_digits(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(pass_code=message.text)
    await call.message.edit_text(text="<b>Введите сумму кредита цифрами:</b>")
    await SG.enter_credit_sum_words.set()



async def enter_credit_sum_words(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(credit_sum_digits=message.text)
    await call.message.edit_text(text="<b>Введите сумму кредита буквами:</b>")
    await SG.enter_cost_sum_digits.set()



async def enter_cost_sum_digits(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(credit_sum_words=message.text)
    await call.message.edit_text(text="<b>Введите стоимость услуги цифрами:</b>")
    await SG.enter_cost_sum_words.set()



async def enter_cost_sum_words(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(cost_sum_digits=message.text)
    await call.message.edit_text(text="<b>Введите стоимость услуги буквами:</b>")
    await SG.enter_percent_digits.set()



async def enter_percent_digits(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(cost_sum_words=message.text)
    await call.message.edit_text(text="<b>Введите процент цифрой:</b>")
    await SG.enter_percent_words.set()



async def enter_percent_words(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(percent_digits=message.text)
    await call.message.edit_text(text="<b>Введите процент буквами:</b>")
    await SG.enter_client_email.set()



async def enter_client_email(message: types.Message, state: FSMContext):
    data = await state.get_data()
    call = data['call']
    await message.delete()
    await state.update_data(percent_words=message.text)
    await call.message.edit_text(text="<b>Введите email клиента:</b>")
    await SG.generate_contract_final.set()


async def generate_contract_final(message: types.Message, state: FSMContext):
    admin_panel_adapter = AdminPanelAdapter()
    keyboard = await admin_panel_adapter.contract_creation_final()

    data = await state.get_data()
    call = data['call']
    await message.delete()
    data = await state.get_data()
    email = message.text
    fio = data['fio']
    pass_issue = data['issuer']
    pass_issue_date = data['pass_issue_date']
    pass_code = data['pass_code']
    passport = data['passport_data']
    birth_date = data['birth_date']
    credit_sum_digits = data['credit_sum_digits']
    credit_sum_words = data['credit_sum_words']
    cost_sum_digits = data['cost_sum_digits']
    cost_sum_words = data['cost_sum_words']
    percent_digits = data['percent_digits']
    percent_words = data['percent_words']
    date = data['date']
    phone_number = data['phone_number']
    date_number = date + "-" + str(random.randint(10000, 99999))


    #print(name, passport_issuer, passportdata, birth_date, credit_sum, date, client_email)





    #generate_contract

    document = Document('files/templates/contract_template.docx')
    for i in document.paragraphs:
        i.text = i.text.replace("(date)", date)
        i.text = i.text.replace("(date-number)", date_number)
        i.text = i.text.replace("(fio)", fio)
        i.text = i.text.replace("(birth_date)", birth_date)
        i.text = i.text.replace("(passport)", passport)
        i.text = i.text.replace("(pass_issue)", pass_issue)
        i.text = i.text.replace("(pass_issue_date)", pass_issue_date)
        i.text = i.text.replace("(pass_code)", pass_code)
        i.text = i.text.replace("(phone_number)", phone_number)
        i.text = i.text.replace("(email)", email)
        i.text = i.text.replace("(cost_sum_digits)", cost_sum_digits)
        i.text = i.text.replace("(cost_sum_words)", cost_sum_words)
        i.text = i.text.replace("(percent)", percent_digits)
        i.text = i.text.replace("(percent_words)", percent_words)
        i.text = i.text.replace("(credit_sum_digits)", credit_sum_digits)
        i.text = i.text.replace("(credit_sum_words)", credit_sum_words)






    await call.message.delete()
    path_to_save = str(pathlib.Path.cwd()) + f"/files/contracts/Договор | {fio}.docx"
    document.tables[-1].cell(1, 1).text = f"""  {fio}                      
Дата рождения: {birth_date}"""
    document.tables[-1].cell(2, 1).text = f"""Паспорт: {passport}                       
Выдан: {pass_issue}       
Дата выдачи: {pass_issue_date}"""
    document.tables[-1].cell(3, 1).text = f"""Телефон: {phone_number} 
 E-mail: {email}"""
    document.save(path_to_save)



    with open(path_to_save, 'rb') as document:
        await call.message.answer_document(caption="<b>Договор сформирован!</b>", document=document, reply_markup=keyboard)













async def send_to_client_btn(call: types.CallbackQuery, state: FSMContext, scheduler: AsyncIOScheduler):
    data = await state.get_data()
    client_id = data.get("client_id")
    user_adapter = UserAdapter()
    admin_panel_adapter = AdminPanelAdapter()
    client_fio = await admin_panel_adapter.get_client_fio_by_id(client_id)
    contract_path = str(pathlib.Path.cwd()) + f"/files/contracts/Договор | {client_fio}.docx"
    await admin_panel_adapter.create_contract(client_id, contract_path)
    user_stage = await user_adapter.get_user_stage(client_id)

    #Если клиент на этапе ожидания договора, то сазу отправляем его
    if user_stage == "prepare_contract":
        contract_aprovement_btn = InlineKeyboardButton('Согласен(-а) ✅', callback_data='contract_approved_btn#do_not_delete_msg')
        back_btn = InlineKeyboardButton('◀️   Назад    ', callback_data='go_back_btn')
        keyboard = InlineKeyboardMarkup().add(contract_aprovement_btn, back_btn)
        await user_adapter.update_user_stage(client_id, "contract_ready", call, scheduler)
        with open(contract_path, "rb") as file:
            await call.message.bot.send_document(caption="<b>Договор готов!</b>\n\nЕсли вы согласны с условиями договора и хотите продолжить работу,"
                                                         " нажмите кнопку для согласия с условиями.", chat_id=client_id, document=file, reply_markup=keyboard)
    #Если клиент на более раннем этапе, то не отправляем ему договор

    await state.finish()
    await call.message.edit_caption(caption="<b>Договор отправлен клиенту, нажмите /start для перехода в меню.</b>")




async def contract_approved_btn(call: types.CallbackQuery, state: FSMContext, scheduler: AsyncIOScheduler):
    user_adapter = UserAdapter()
    admin_panel_adapter = AdminPanelAdapter()
    await admin_panel_adapter.make_contract_approved(call.from_user.id)
    await user_adapter.update_user_stage(call.from_user.id, "check_credit_history", call, scheduler)

    base_adapter = BaseAdapter()
    keyboard = await base_adapter.start(int(call.from_user.id))
    await call.message.answer(text=base_adapter._msg, reply_markup=keyboard)
    if call.data.split("#")[-1] == "delete_msg":
        await call.message.delete()
    else:
        await call.message.edit_caption(caption="<b>Условия договора приняты ✅</b>")









# pagination handler
async def contract_clients_list_pagination(call: types.CallbackQuery, state: FSMContext):
    paginator = (await state.get_data()).get('paginator')
    keyboard = paginator.page_switch(call.data.split('#')[1])
    try:
        await call.message.edit_text(text="<b>Выберите клиента:</b>", reply_markup=keyboard)
        await state.update_data(paginator=paginator)
    except exceptions.MessageNotModified:
        await call.answer(text='Больше нет клиентов!', show_alert=True)


def register_generate_contract_handlers(dp: Dispatcher):
    dp.register_message_handler(
        enter_birth_date,
        state=SG.enter_birth_date
    )
    dp.register_message_handler(
        enter_passport_data,
        state=SG.enter_passport_data
    )
    dp.register_message_handler(
        enter_passport_issue_date,
        state=SG.enter_passport_issue_date
    )
    dp.register_message_handler(
        enter_passport_issue_code,
        state=SG.enter_passport_issue_code
    )
    dp.register_message_handler(
        enter_date,
        state=SG.enter_date
    )
    dp.register_message_handler(
        enter_passport_issuer,
        state=SG.enter_passport_issuer
    )
    dp.register_message_handler(
        enter_credit_sum_digits,
        state=SG.enter_credit_sum
    )
    dp.register_message_handler(
        enter_credit_sum_words,
        state=SG.enter_credit_sum_words
    )
    dp.register_message_handler(
        enter_cost_sum_digits,
        state=SG.enter_cost_sum_digits
    )
    dp.register_message_handler(
        enter_cost_sum_words,
        state=SG.enter_cost_sum_words
    )
    dp.register_message_handler(
        enter_percent_digits,
        state=SG.enter_percent_digits
    )
    dp.register_message_handler(
        enter_percent_words,
        state=SG.enter_percent_words
    )
    dp.register_message_handler(
        enter_client_email,
        state=SG.enter_client_email
    )
    dp.register_message_handler(
        generate_contract_final,
        state=SG.generate_contract_final
    )
    dp.register_callback_query_handler(
        send_to_client_btn,
        lambda callback_query: callback_query.data == 'send_to_client_btn',
        state='*'
    )
    dp.register_callback_query_handler(
        choose_client_to_create_contract,
        lambda callback_query: callback_query.data == "generate_contract_btn",
        state='*'
    )
    dp.register_callback_query_handler(
        contract_clients_list_pagination,
        lambda callback_query: callback_query.data.split("#")[0] == 'amg',
        state=SG.clients_waiting_contract_pagination
    )
    dp.register_callback_query_handler(
        enter_birth_date,
        lambda callback_query: callback_query.data.split("#")[1] == "user_number",
        state=SG.clients_waiting_contract_pagination,
    )
    dp.register_callback_query_handler(
        contract_approved_btn,
        lambda callback_query: callback_query.data.split("#")[0] == 'contract_approved_btn',
        state='*'
    )
    dp.register_callback_query_handler(
        get_contract_btn,
        lambda callback_query: callback_query.data == 'get_contract_btn',
        state='*'
    )